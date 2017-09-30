import os
import re
import wx
import xlrd
import xlwt
from openpyxl import load_workbook
from xlutils.copy import copy
from xlutils.styles import Styles
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Cm
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import frame

COLUMN_QUESTION_WIDTH = 2
COLUMN_TOPIC_WIDTH = 4
COLUMN_MARKER_WIDTH = 3
COLUMN_COMMENTS_WIDTH = 8
GRAY = 'BBBBBB'

subjects_all = {}
subject_map_workbook = None
students_all = []
mcq_response_percentage_highlight_threshold = 0
mcq_question_length = 0
mcq_responses_all = {}
mcq_question_numbers = []
mcq_question_answers = []
mcq_question_weightages = []
non_mcq_question_length = 0
non_mcq_question_names = []
non_mcq_question_weightages = []
non_mcq_results_all = []
non_mcq_questions_path = ''
valid_mcq_responses_selected = False
valid_mcq_answers_selected = False
valid_non_mcq_questions_selected = False
include_mcq = True
include_non_mcq = True

class MainFrame(frame.Frame):
        def error_dialog(self, message):
                wx.MessageBox(message, 'Error', wx.OK|wx.ICON_ERROR)

        def subject_mapping_changed(self, event):
                global subject_map_workbook
                path = event.GetPath()
                try:
                        subject_map_workbook = xlrd.open_workbook(path)
                except xlrd.biffh.XLRDError:
                        self.error_dialog('Invalid file')
                        return
                for sheet_num, sheet in enumerate(subject_map_workbook.sheets()):
                        batch_subjects = set()
                        for subject in sheet.col_values(6):
                                # We do not want the header
                                if subject != 'SUBJECT CODE':
                                        batch_subjects.add(subject)
                        sheet_name = subject_map_workbook.sheet_names()[sheet_num]
                        subjects_all.update({sheet_name: batch_subjects})       
                batches_sorted = sorted(list(subjects_all.keys()))
                self.listbox_batch.Set(sorted(list(batches_sorted)))

        def batch_changed(self, event):
                batch = event.GetString()
                subjects_sorted = sorted(list(subjects_all[batch]))
                self.listbox_subject.Set(subjects_sorted)

        def students_by_subject_code(self, subject_code, sheet):
                '''Retrieve all the students taking a subject with a subject code'''
                students = []
                student_no = 1
                for student in sheet.get_rows():
                        if student[6].value == subject_code:
                                student_class = student[3].value
                                student_register_number = student[2].value
                                student_nric = student[0].value
                                student_name = student[1].value
                                student_teacher = student[7].value
                                students.append([student_no, student_class, student_register_number,
                                                                 student_nric, student_name, student_teacher])
                                student_no += 1
                return students

        def get_non_mcq_questions(self):
                '''Read self.grid_non_mcq_questions for non-MCQ question entries'''
                non_mcq_questions = []
                row_num = 0
                while True:
                        non_mcq_question_name = self.grid_non_mcq_questions.GetCellValue(row_num, 0)
                        non_mcq_question_full_mark = self.grid_non_mcq_questions.GetCellValue(row_num, 1)
                        if non_mcq_question_name == '' and non_mcq_question_full_mark == '':
                                # Treat empty entry as end of data input
                                break
                        elif non_mcq_question_name == '' or non_mcq_question_full_mark == '':
                                # One of the fields is missing
                                self.error_dialog('Row {}: missing field'.format(row_num+1))
                                return
                        else:
                                try:
                                        # Validate that full marks is a number
                                        int(non_mcq_question_full_mark)
                                except ValueError:
                                        self.error_dialog('Row {}: full marks must be a number'.format(row+1))
                                        return
                                non_mcq_questions.append([non_mcq_question_name, non_mcq_question_full_mark])
                        row_num += 1
                return non_mcq_questions

        def generate_template(self, event):
                wb = load_workbook("template.xlsx")
                ws = wb.active
                msg = "Please wait while we process generate the template file..."
                busyDlg = wx.BusyInfo(msg)                
                try:
                        batch = self.listbox_batch.GetString(self.listbox_batch.GetSelection())
                        subject_code = self.listbox_subject.GetString(self.listbox_subject.GetSelection())
                except wx._core.PyAssertionError:
                        self.error_dialog('Please select a batch and subject first')
                        return
                subjects_sheet = subject_map_workbook.sheet_by_name(batch)
                students = self.students_by_subject_code(subject_code, subjects_sheet)

                # Write student data for the particular subject
                
                try:
                    for sidx,s in enumerate(students):
                            for cidx, info in enumerate(s):
                                    ws.cell(column=cidx+1, row=sidx+3, value=info)
                    offset = len(students[0])
                    for sidx, q in enumerate(self.get_non_mcq_questions()):
                            for cidx, info in enumerate(q):
                                    ws.cell(column=sidx+offset+1, row=1+cidx, value=info)
                                    
                    wb.save(filename = '{}-{}.xls'.format(batch, subject_code))
                except Exception as e:
                    self.error_dialog('Error: {}'.format(e))
                    busyDlg = None
                    raise e
                else:
                    pass
                # Write non-MCQ question data into Excel file
                busyDlg = None

                
        def clear_non_mcq_questions(self, event):
                for row_num in range(self.grid_non_mcq_questions.GetNumberRows()):
                        for column_num in range(self.grid_non_mcq_questions.GetNumberCols()):
                                self.grid_non_mcq_questions.SetCellValue(row_num, column_num, '')

        def include_mcq_toggled(self, event):
                global include_mcq
                
                if event.GetEventObject().GetValue():
                        self.text_mcq_responses_dir.Enable()
                        self.dirpicker_mcq_responses.Enable()
                        self.text_mcq_answers.Enable()
                        self.filepicker_mcq_answers.Enable()
                        include_mcq = True
                else:
                        self.text_mcq_responses_dir.Disable()
                        self.dirpicker_mcq_responses.Disable()
                        self.text_mcq_answers.Disable()
                        self.filepicker_mcq_answers.Disable()
                        include_mcq = False
                        mcq_question_length = 0
                

        def include_non_mcq_toggled(self, event):
                global include_non_mcq
                if event.GetEventObject().GetValue():
                        non_mcq_question_length = 0
                        include_non_mcq = True
                        #self.text_non_mcq_questions.Enable()
                        #self.filepicker_non_mcq_questions.Enable()
                else:
                        include_non_mcq = False
                        #self.text_non_mcq_questions.Disable()
                        #self.filepicker_non_mcq_questions.Disable()
                
        def mcq_dir_changed(self, event):
                global valid_mcq_responses_selected, mcq_responses_all
                
                mcq_path = event.GetPath()
                mcq_answers_all_path = mcq_path + '/scoring.xlsx'
                # The regex basically checks if the filename starts with "Student Responses" and is of type .dif
                mcq_responses_filenames = ['{}/{}'.format(mcq_path, filename) for filename in os.listdir(mcq_path)
                                                                   if re.match('^Student Responses - .*\.dif$', filename) != None]
                if mcq_responses_filenames == []:
                        valid_mcq_responses_selected = False
                        self.error_dialog('No files containing student responses found. ' + 
                                                          'Make sure they start with "Student Responses".')
                        return
                valid_mcq_responses_selected = True
                for filename in mcq_responses_filenames:
                        with open(filename, 'rb') as mcq_responses_file:
                                lines = mcq_responses_file.readlines()
                                lines = [line.strip('\r\n') for line in lines]
                                # The line goes "Class : <Batch>-<Class Name>" so get only the last 2 characters
                                class_name = lines[0].strip('\t\t')[-2:]
                                # Ignore the metadata at the front as well as the heading
                                mcq_responses_raw = [line.split('\t') for line in lines[6:]]
                                # student_responses[0] gives the register number
                                # Convert each row to a tuple as dictionary keys are immutable unlike lists
                                # The (class_name, register) will be used to index the students' responses
                                students_class_register_no = \
                                                [tuple([class_name]+[int(student_responses[0])]) for student_responses in mcq_responses_raw]
                                # [3:-3] includes just the actual responses
                                mcq_responses_all_questions = [student_responses[3:-3] for student_responses in mcq_responses_raw]
                                mcq_responses_all.update(dict(zip(students_class_register_no, mcq_responses_all_questions)))
        
        def mcq_answer_template_changed(self, event):
                global valid_mcq_answers_selected, mcq_question_answers, mcq_question_weightages, \
                           mcq_question_length, mcq_question_numbers
                mcq_answers_path = event.GetPath()
                try:
                        #mcq_df = pd.read_excel(mcq_answers_path)
                        mcq_df = pd.read_csv(mcq_answers_path, delimiter='\t')
                except xlrd.biffh.XLRDError:
                        valid_mcq_answers_selected = False
                        self.error_dialog('Invalid file')
                        return
                valid_mcq_answers_selected = True
                
                
                
##                # [1:] ignores the vertical header
##                mcq_question_length = len(mcq_df.columns[1:])
##                # pandas reads the first row of the table as the column labels
##                mcq_question_numbers = mcq_df.columns[1:].tolist()
##                mcq_question_answers = mcq_df.iloc[0, 1:].tolist()
##                mcq_question_weightages = pd.to_numeric(mcq_df.iloc[1, 1:]).tolist()
                mcq_question_length = len(mcq_df) - 4
                mcq_question_numbers = mcq_df.iloc[4:,0].tolist()
                mcq_question_answers = mcq_df.iloc[4:,1].tolist()
                mcq_question_weightages = pd.to_numeric(mcq_df.iloc[4:,2]).tolist()
                
        def non_mcq_questions_changed(self, event):
                global non_mcq_questions_path, valid_non_mcq_questions_selected, students_all, \
                           non_mcq_question_names, non_mcq_question_weightages, non_mcq_question_length, \
                           non_mcq_results_all
                non_mcq_questions_path = event.GetPath()
                try:
                        non_mcq_df = pd.read_excel(non_mcq_questions_path)
                except xlrd.biffh.XLRDError:
                        valid_non_mcq_questions_selected = False
                        self.error_dialog('Invalid file')
                        return
                valid_non_mcq_questions_selected = True
                # [1:] ignores the full marks row, [1:5] ignores the non-MCQ results
                # df.values.tolist() converts the DataFrame into a list of lists
                students_all = non_mcq_df.iloc[1:, 1:5].values.tolist()
                # Hence, [6:] are the actual non-MCQ results
                non_mcq_question_length = len(non_mcq_df.columns[6:])
                # pandas reads the first row of the table as the column labels
                non_mcq_question_names = non_mcq_df.columns[6:].tolist()
                # [0] selects the full marks row
                # Call pd.to_numeric() as the numerical value is needed to calculate FI and DI
                non_mcq_question_weightages = pd.to_numeric(non_mcq_df.iloc[0, 6:]).tolist()
                non_mcq_results_all = non_mcq_df.iloc[1:, 6:].values.tolist()

        def valid_input(self):
                '''
                Ensure that valid files and directories have been selected in the 2nd tab
                Returns False if not all valid files and directories have been picked, True otherwise
                '''
                if not include_mcq and not include_non_mcq:
                        self.error_dialog('Please include MCQ, non-MCQ, or both types of items first')
                        return False

                if include_mcq and not valid_mcq_responses_selected:
                        self.error_dialog('Please select a valid MCQ responses directory first')
                        return False
                if include_mcq and not valid_mcq_answers_selected:
                        self.error_dialog('Please select a valid MCQ answer file first')
                        return False
                # Do not check value of include_non_mcq as we need data from the non-MCQ file no matter what
                if not valid_non_mcq_questions_selected:
                        self.error_dialog('Please select a valid student list first')
                        return False

                try:
                        threshold = int(self.input_mcq_reponse_percentage_highlight_threshold.GetValue())
                except ValueError:
                        self.error_dialog('Threshold must be an integer')
                        return False
                if threshold <= 0 or threshold >= 100:
                        self.error_dialog('Threshold must be between 1 and 100')
                        return False
                # Sorry, couldn't think of a shorter name
                global mcq_response_percentage_highlight_threshold
                mcq_response_percentage_highlight_threshold = threshold
                return True

        def merge_and_analyse(self, event):
                if not self.valid_input():
                        return
                msg = "Please wait while we process your request..."
                busyDlg = wx.BusyInfo(msg)
                try:
                        template_workbook = xlrd.open_workbook(non_mcq_questions_path, formatting_info=False)
                except Exception as e:
                        self.error_dialog('Error: {}'.format(e))
                        busyDlg = None
                        return
                student_list_workbook = copy(template_workbook)
                student_list_sheet = student_list_workbook.get_sheet(0)
                # If only include non-MCQ then the resulting student list will be exactly the same, hence save as a direct copy
                if not (include_non_mcq and not include_mcq):
                        if include_mcq:
                                self.write_rows(0, 6, [mcq_question_numbers, mcq_question_weightages], student_list_sheet, styled=False)
                                for index, student in enumerate(students_all):
                                        # [:2] includes just the class and register number, which are the fields needed to index the responses
                                        student = tuple(student[:2])
                                        try:
                                                mcq_responses_all[student]
                                        except KeyError:
                                                # The student did not sit for the MCQ test
                                                self.write_column(index+2, 6, ['NA']*mcq_question_length, student_list_sheet, styled=False)
                                        else:
                                                for index_inner, column_num in enumerate(range(6, 6+mcq_question_length)):
                                                        mark = (mcq_question_weightages[index_inner] 
                                                                        if mcq_responses_all[student][index_inner] == mcq_question_answers[index_inner]
                                                                        else 0)
                                                        student_list_sheet.write(index+2, column_num, mark)
                        if include_non_mcq:
                                self.write_rows(0, 6+mcq_question_length, [non_mcq_question_names, non_mcq_question_weightages],
                                                                student_list_sheet, styled=False)
                                self.write_rows(2, 6+mcq_question_length, non_mcq_results_all, student_list_sheet, styled=False)

                # Ignore the preceding path as well as the student extension
                #print(non_mcq_questions_path)
                non_mcq_questions_filename_short = non_mcq_questions_path.split('\\')[-1].split('.')[0]
                student_list_filename_map = {include_mcq and not include_non_mcq: '{}_mcq.xlsx'.format(non_mcq_questions_filename_short),
                                                                 include_non_mcq and not include_mcq: '{}_nonmcq.xlsx'.format(non_mcq_questions_filename_short),
                                                                         include_mcq and include_non_mcq: '{}_all.xls'.format(non_mcq_questions_filename_short)}
                print(student_list_filename_map)
                try:
                        # Make a new directory and save all the generated .docx and .xlsx files there
                        #print(non_mcq_questions_filename_short)
                        os.mkdir(non_mcq_questions_filename_short)
                except OSError:
                        # Directory already exists
                        pass
                filepath = os.path.dirname( __file__ ) #remember current directory
                os.chdir(non_mcq_questions_filename_short)
                student_list_workbook.save(student_list_filename_map[True])

                # Used for generating analysis for each group (classes + batch)
                # set() removes duplicate values
                group_names = ['batch'] + list(set([student[0] for student in students_all]))
                mcq_responses_list = []
                if include_mcq:
                        for student_class_register_number, student_responses in mcq_responses_all.iteritems():
                                # Find the corresponding name to include in the data report
                                for student in students_all:
                                        # [:2] is the class and register number
                                        if student[:2] == list(student_class_register_number):
                                                student_name = student[3]
                                mcq_responses_list.append(list(student_class_register_number) + \
                                                                                  [student_name] + \
                                                                                  student_responses)
                        mcq_responses_df = pd.DataFrame(mcq_responses_list,
                                                                                        columns = ['Class', 'Class Reg No', 'Name'] + mcq_question_numbers)
                        mcq_grouped = mcq_responses_df.groupby('Class', axis=0)
                        # Master list of all MCQ option percentages
                        mcq_percentages_df = pd.DataFrame(columns=mcq_question_numbers)
                        mcq_students_all = {}
                        # [3:] only includes the actual responses
                        for class_name, responses in mcq_grouped:
                                mcq_percentage_class, mcq_students = self.process_mcq_responses(class_name, responses)
                                mcq_percentages_df = mcq_percentages_df.append(mcq_percentage_class)
                                mcq_students_all.update({class_name: mcq_students})
                        mcq_percentage_batch, mcq_students = self.process_mcq_responses('batch', mcq_responses_df)
                        mcq_percentages_df = mcq_percentages_df.append(mcq_percentage_batch)
                        mcq_students_all.update({'batch': mcq_students})

                # Get FIDI data
                results_df = pd.read_excel(student_list_filename_map[True])
                # Remove headers and student info
                results_df = results_df.drop(results_df.index[:1], axis=0)
                # Class, register number and name needed for data drill down
                results_df = results_df.drop(['SN', 'NRIC', 'Subject Teacher'], axis=1)
                # Exclude all students who didn't take the test from FIDI calculations
                results_df.dropna()
                # Needed to find top and bottom scorers
                # 3: to ignore class, register number and name
                results_df['total_score'] = pd.Series(results_df.iloc[:, 3:].sum(axis=1))
                results_grouped = results_df.groupby('Class', axis=0)

                # Dynamic list based on whether MCQ, non-MCQ or both types of questions are included
                weightage_list = []
                if include_mcq:
                        weightage_list += mcq_question_weightages
                if include_non_mcq:
                        weightage_list += non_mcq_question_weightages
                mcq_questions_name = range(1, mcq_question_length+1)
                fidi_df = pd.DataFrame()
                # Process individual class data
                top_students = {}
                bottom_students = {}
                for class_name, results in results_grouped:
                        fi_df, di_df, class_top_students, class_bottom_students = self.process_fidi(class_name, results, weightage_list)
                        for df in (fi_df, di_df):
                                fidi_df = fidi_df.append(df)
                        top_students.update({class_name: class_top_students})
                        bottom_students.update({class_name: class_bottom_students})
                # Process batch data
                fi_df, di_df, batch_top_students, batch_bottom_students = self.process_fidi('batch', results_df, weightage_list)
                for df in (fi_df, di_df):
                        fidi_df = fidi_df.append(df)
                top_students.update({'batch': batch_top_students})
                bottom_students.update({'batch': batch_bottom_students})
                # Conveniently, if only non-MCQ questions are selected, then mcq_questions_length will be 0,
                # hence mcq_fidi_df will be empty and non_mcq_fidi_df will have the expected values
                mcq_fidi_df = fidi_df.iloc[:, :mcq_question_length]
                non_mcq_fidi_df = fidi_df.iloc[:, mcq_question_length:]

                if include_non_mcq:
                        below_fi_students_all = self.process_below_fi_students(results_df, non_mcq_fidi_df, group_names)

                # Write FIDI Analysis to Excel files for each group
                # Default green is too dark
                xlwt.add_palette_colour('light_dark_green', 21)
                fi_style_map = {lambda x: x == 0: self.cell_style(),
                                                lambda x: 0 < x < 0.4: self.cell_style(bg='red'),
                                                lambda x: 0.4 <= x <= 0.69: self.cell_style(bg='light_orange'),
                                                lambda x: 0.7 <= x < 1: self.cell_style(bg='light_blue'),
                                                lambda x: x == 1: self.cell_style()}
                di_style_map = {lambda x: -1 <= x <= 0: self.cell_style(),
                                                lambda x: 0 < x <= 0.19: self.cell_style(bg='light_orange'),
                                                lambda x: 0.2 <= x <= 0.29: self.cell_style(bg='light_green'),
                                                lambda x: 0.3 <= x < 1: self.cell_style(bg='light_dark_green'),
                                                lambda x: x == 1: self.cell_style()}
                mcq_style_map = {lambda x: x < mcq_response_percentage_highlight_threshold: self.cell_style(fg='red'),
                                                 lambda x: x >= mcq_response_percentage_highlight_threshold: self.cell_style()}

                mcq_headings_class = [('Marks', self.cell_style(bg='yellow', bold=True)),
                                                          ('Qn No.', self.cell_style(bg='yellow', bold=True)),
                                                          ('Class FI', self.cell_style(bg='light_green', bold=True)),
                                                          ('Class DI', self.cell_style(bg='light_dark_green', bold=True)),
                                                          ('Cohort FI', self.cell_style(bg='light_green', bold=True)),
                                                          ('Cohort DI', self.cell_style(bg='light_dark_green', bold=True)),
                                                          ('A', self.cell_style(bold=True)),
                                                          ('B', self.cell_style(bold=True)),
                                                          ('C', self.cell_style(bold=True)),
                                                          ('D', self.cell_style(bold=True))]
                mcq_headings_cohort_unwanted = ['Class FI', 'Class DI']
                mcq_headings_cohort = [(heading, style) for heading, style in mcq_headings_class 
                                                           if heading not in mcq_headings_cohort_unwanted]

                non_mcq_headings_class = [('Marks', self.cell_style(bg='yellow', bold=True)),
                                                                  ('Qn No.', self.cell_style(bg='yellow', bold=True)),
                                                                  ('Class FI', self.cell_style(bg='light_green', bold=True)),
                                                                  ('Class DI', self.cell_style(bg='light_dark_green', bold=True)),
                                                                  ('Cohort FI', self.cell_style(bg='light_green', bold=True)),
                                                                  ('Cohort DI', self.cell_style(bg='light_dark_green', bold=True))]
                non_mcq_headings_cohort_unwanted = ['Class FI', 'Class DI']
                non_mcq_headings_cohort = [(heading, style) for heading, style in non_mcq_headings_class
                                                                   if heading not in non_mcq_headings_cohort_unwanted]

                top_bottom_students_table_header_style = [self.cell_style(bg='yellow'),
                                                                                                  self.cell_style(bg='yellow'),
                                                                                                  self.cell_style(bg='white'),
                                                                                                  self.cell_style(bg='white')]
                top_bottom_students_table_header = ['Class', 'Reg No', 'Name', 'Score']
                top_bottom_students_headers = ['Top 27%', 'Bottom 27%']

                mcq_students_table_header = top_bottom_students_table_header[:-1]
                mcq_students_table_header_style = top_bottom_students_table_header_style[:-1]

                below_fi_students_table_header = mcq_students_table_header
                below_fi_students_table_header_style = mcq_students_table_header_style

                # Get data from the whole batch
                if include_mcq:
                        mcq_question_info_styled = self.get_analysis_question_info_styled(mcq_question_weightages,
                                                                                                                                                          mcq_question_numbers,
                                                                                                                                                          mcq_question_length)
                        mcq_percentages_batch_styled = self.get_analysis_mcq_percentages_styled(mcq_percentages_df,
                                                                                                                                                                        mcq_style_map,
                                                                                                                                                                        'batch')
                        mcq_fidi_batch_styled = self.get_analysis_fidi_styled(mcq_fidi_df, fi_style_map, di_style_map, 'batch')
                if include_non_mcq:
                        non_mcq_question_info_styled = self.get_analysis_question_info_styled(non_mcq_question_weightages,
                                                                                                                                                                  non_mcq_question_names,
                                                                                                                                                                  non_mcq_question_length)
                        non_mcq_fidi_batch_styled = self.get_analysis_fidi_styled(non_mcq_fidi_df, fi_style_map, di_style_map, 'batch')

                for group_name in group_names:
                        # Get data from each group
                        analysis_workbook = xlwt.Workbook()
                        analysis_workbook.set_colour_RGB(21, 76, 188, 49)
                        fidi_sheet = analysis_workbook.add_sheet('FIDI Analysis')
                        top_bottom_sheet = analysis_workbook.add_sheet('Top27Bottom27')
                        fidi_row_num = fidi_column_num = 0
                        analysis_name_map = {include_mcq and not include_non_mcq:
                                                                        '{}_mcq_analysis.xls'.format(group_name),
                                                                 not include_mcq and include_non_mcq:
                                                                        '{}_nonmcq_analysis.xls'.format(group_name),
                                                                 include_mcq and include_non_mcq:
                                                                        '{}_all_analysis.xls'.format(group_name)}

                        if include_mcq:
                                mcq_students_sheet = analysis_workbook.add_sheet('MCQ Results')
                                mcq_students_cells_all = self.write_mcq_students(mcq_question_answers,
                                                                                                                                 mcq_students_all[group_name],
                                                                                                                                 mcq_students_table_header,
                                                                                                                                 mcq_students_table_header_style,
                                                                                                                                 mcq_students_sheet)
                                mcq_fidi_group_styled = self.get_analysis_fidi_styled(mcq_fidi_df, fi_style_map,
                                                                                                                                          di_style_map, group_name)
                                mcq_percentages_group_styled = self.get_analysis_mcq_percentages_styled(mcq_percentages_df,
                                                                                                                                                                                 mcq_style_map,
                                                                                                                                                                                 group_name)

                                # Insert hyperlinks
                                for row_num, row in enumerate(mcq_percentages_group_styled):
                                        for item_num, item in enumerate(row):
                                                # Regarding mcq_students_cells_all[row_num][item_num][0] we can operate under the assumption
                                                # that the number of columns won't exceed 26
                                                mcq_percentages_group_styled[row_num][item_num][0] = \
                                xlwt.Formula('HYPERLINK("[{}]\'MCQ Results\'!{}{}";"{}")'.format(analysis_name_map[True],
                                                                                                                                                                 mcq_students_cells_all[row_num][item_num][0],
                                                                                                                                                                 mcq_students_cells_all[row_num][item_num][1:],
                                                                                                                                                                 item[0]))

                                mcq_table = []
                                if group_name == 'batch':
                                        mcq_table_no_heading = mcq_question_info_styled + mcq_fidi_group_styled + \
                                                                                   mcq_percentages_group_styled
                                        for index in range(len(mcq_table_no_heading)):
                                                mcq_table.append([mcq_headings_cohort[index]] + mcq_table_no_heading[index])
                                else:
                                        # Include cohort FIDI when processing the classes but not the whole cohort
                                        mcq_table_no_heading = mcq_question_info_styled + mcq_fidi_group_styled + \
                                                                                   mcq_fidi_batch_styled + mcq_percentages_group_styled 
                                        for index in range(len(mcq_table_no_heading)):
                                                mcq_table.append([mcq_headings_class[index]] + mcq_table_no_heading[index])

                                fidi_sheet.write(0, 0, 'MCQ - FIDI with Distractor Analysis')
                                mcq_table_header = [row[0] for row in mcq_table]
                                # Headers are not needed for chunks
                                mcq_table = [row[1:] for row in mcq_table]
                                mcq_table_chunks = self.create_chunks(16, mcq_table)
                                fidi_row_num = 1
                                fidi_row_num = self.write_chunks(fidi_row_num, 0, mcq_table_header, mcq_table_chunks, fidi_sheet)

                        if include_non_mcq:
                                below_fi_students_sheet = analysis_workbook.add_sheet('Students below FI')
                                below_fi_students_cells_all = self.write_below_fi_students(below_fi_students_all[group_name],
                                                                                                                                                          below_fi_students_table_header,
                                                                                                                                                          below_fi_students_table_header_style,
                                                                                                                                                          below_fi_students_sheet)
                                non_mcq_fidi_group_styled = self.get_analysis_fidi_styled(non_mcq_fidi_df, fi_style_map,
                                                                                                                                                  di_style_map, group_name)
                                # Insert hyperlinks
                                for item_num, item in enumerate(non_mcq_fidi_group_styled[0]):
                                        non_mcq_fidi_group_styled[0][item_num][0] = \
                                xlwt.Formula('HYPERLINK("[{}]\'Students below FI\'!{}{}","{}")'.format(analysis_name_map[True],
                                                                                                                                                                           below_fi_students_cells_all[item_num][0],
                                                                                                                                                                           below_fi_students_cells_all[item_num][1:],
                                                                                                                                                                           item[0]))

                                non_mcq_table = []
                                if group_name == 'batch':
                                        non_mcq_table_no_heading = non_mcq_question_info_styled + non_mcq_fidi_group_styled
                                        for index in range(len(non_mcq_table_no_heading)):
                                                non_mcq_table.append([non_mcq_headings_cohort[index]] + non_mcq_table_no_heading[index])
                                else:
                                        non_mcq_table_no_heading = non_mcq_question_info_styled + non_mcq_fidi_group_styled + \
                                                                                           non_mcq_fidi_batch_styled
                                        for index in range(len(non_mcq_table_no_heading)):
                                                non_mcq_table.append([non_mcq_headings_class[index]] + non_mcq_table_no_heading[index])

                                fidi_sheet.write(fidi_row_num, 0, 'Non-MCQ - FIDI')
                                non_mcq_table_header = [row[0] for row in non_mcq_table]
                                non_mcq_table = [row[1:] for row in non_mcq_table]
                                non_mcq_table_chunks = self.create_chunks(16, non_mcq_table)
                                fidi_row_num += 1
                                fidi_row_num = self.write_chunks(fidi_row_num, 0, non_mcq_table_header, non_mcq_table_chunks, fidi_sheet)

                        interpretation_note_style = self.cell_style(fg='red', bold=True, wrap=True)
                        interpretation_note_style = xlwt.easyxf('font: colour red, bold on; align: wrap on, vert center')
                        interpretation_note_paragraph = \
                                        "Note: Interpretations provided below serve as a guide to help surface items for further discussion, " + \
                                        "analysis and any possible revision of items. Please bear in mind that the range of values given by " + \
                                        "the reseachers were based on dichotomous test items. It is essential to cross reference the FI&DI " + \
                                        "values with TOS and the assessment purpose of the individual item when reviewing all test items and " + \
                                        "preparing markers' report"
                        fidi_sheet.row(fidi_row_num).height_mismatch = True
                        fidi_sheet.row(fidi_row_num).height = 2000
                        fidi_sheet.write_merge(fidi_row_num, fidi_row_num, 0, 10,
                                                                   interpretation_note_paragraph, interpretation_note_style)
                        fidi_row_num += 2
                        self.write_fidi_interpretation_table(fidi_row_num, fidi_column_num, fidi_sheet)
                        self.write_top_bottom_students(top_students, bottom_students, top_bottom_students_table_header,
                                                                                   top_bottom_students_table_header_style, top_bottom_students_headers,
                                                                                   group_name, top_bottom_sheet)

                        analysis_workbook.save('{}'.format(analysis_name_map[True]))

                # Generate marker's report
                report = Document()
                report.add_heading("Markers' Report ({})".format(non_mcq_questions_filename_short), 0)
                report.add_heading("Item Analysis FIDI", 1)
                if include_mcq:
                        mcq_fi_all = [item[0] for item in mcq_fidi_batch_styled[0]]
                        mcq_di_all = [item[0] for item in mcq_fidi_batch_styled[1]]
                        mcq_easy, mcq_moderate, mcq_difficult = self.calculate_questions_difficulty(mcq_fi_all)
                        report.add_heading(
                                        'MCQ - {}% easy questions, {}% moderately difficult questions, {}% difficult questions'
                                                .format(mcq_easy, mcq_moderate, mcq_difficult), 2
                        )
                if include_non_mcq:
                        non_mcq_fi_all = [item[0] for item in non_mcq_fidi_batch_styled[0]]
                        non_mcq_di_all = [item[0] for item in non_mcq_fidi_batch_styled[1]]
                        non_mcq_easy, non_mcq_moderate, non_mcq_difficult = self.calculate_questions_difficulty(non_mcq_fi_all)
                        report.add_heading(
                                        'Non MCQ - {}% easy questions, {}% moderately difficult questions, {}% difficult questions'
                                                .format(non_mcq_easy, non_mcq_moderate, non_mcq_difficult), 2
                        )
                spacer = report.add_paragraph()
                spacer_run = spacer.add_run()
                spacer_run.add_break()
                if include_mcq:
                        self.create_section_template('MCQ', [str(i+1) for i in range(mcq_question_length)],
                                                                                 mcq_fi_all, mcq_di_all, report)
                        spacer = report.add_paragraph()
                        spacer_run = spacer.add_run()
                        spacer_run.add_break()
                if include_non_mcq:
                        self.create_section_template('Non MCQ', non_mcq_question_names, 
                                                                                 non_mcq_fi_all, non_mcq_di_all, report)
                report_filename_map = {include_mcq and not include_non_mcq: 'markers_report_mcq.docx',
                                                          include_non_mcq and not include_mcq: 'markers_report_non_mcq.docx',
                                                          include_mcq and include_non_mcq: 'markers_report_all.docx'}
                report.save(report_filename_map[True])
                os.chdir(os.path.join( os.path.dirname( __file__ ), '..' )) #change directory back to curr directory
                busyDlg = None

        def get_analysis_question_info_styled(self, question_weightages, question_names, question_length):
                '''Retrieve the styled question info list to draw in the table'''
                question_info_style = [self.cell_style(bg='yellow')] * question_length
                question_info_styled = [zip(item, question_info_style) for item in (question_weightages, question_names)]
                return question_info_styled
        
        def get_analysis_mcq_percentages_styled(self, option_percentages_df, style_map, group_name):
                '''Retrieve the styled MCQ option percentage info to draw in the table'''
                option_percentages_indices = ['{}_{}'.format(group_name, option) for option in 'ABCD']
                option_percentages = [option_percentages_df.loc[index].tolist() for index in option_percentages_indices]
                # Use lists as the value of the item will be changed later to include the hyperlink
                option_percentages_styled = [[[item, self.corresponding_style(style_map, item)]
                                                                          for item in option_percentages[i]]
                                                                          for i in range(4)]    
                return option_percentages_styled

        def get_analysis_fidi_styled(self, fidi_df, fi_style_map, di_style_map, group_name):
                '''Retrieve the styled FIDI list to draw in the table'''
                fidi_style_maps = (fi_style_map, di_style_map)
                fidi_indices = ['{}_fi'.format(group_name), '{}_di'.format(group_name)]
                fidis = [fidi_df.loc[index].tolist() for index in fidi_indices]
                # Refer to comment in above function
                fidi_styled = [[[item, self.corresponding_style(fidi_style_maps[i], item)]
                                                for item in fidis[i]]
                                                for i in range(2)]
                return fidi_styled

        def create_chunks(self, chunk_size, data):
                '''Divide data into smaller chunks that will be displayed one below the other'''
                chunks = []
                while True:
                        if len(data[0]) >= chunk_size:
                                # Chunks can still be created
                                chunks.append([row[:chunk_size] for row in data])
                                # Remove the columns from the original data as it is already in the chunk
                                data = [row[chunk_size:] for row in data]
                        else:
                                # No chunks need to be formed
                                chunks.append(data)
                                break
                return chunks
        
        def write_chunks(self, row_num, column_num, header, chunks, sheet, columns_length=None):
                '''Write chunks to sheet'''
                for chunk in chunks:
                        # Ignore the row_num as we will be writing the question values in the same row
                        self.write_column(row_num, column_num, header, sheet)
                        row_num = self.write_rows(row_num, column_num+1, chunk, sheet, columns_length=columns_length)
                        row_num += 2
                return row_num

        def write_fidi_interpretation_table(self, row_num, column_num, sheet):
                '''Populate tables of FI and DI interpretations for reference'''
                fi_values_list = ['FI', '0', '< 0.40', '0.4-0.69', '>= 0.70', '1']
                fi_interpretation_list = ['Interpretation', 'Very Difficult Item', 'Difficult Item', 'Moderately Difficult Item',
                                                                  'Easy Item', 'Very Easy Item']
                fi_colour_list = ['light_green', 'white', 'red', 'light_orange', 'light_blue', 'white'] 
                fi_values_list = zip(fi_values_list, [self.cell_style(bg=colour) for colour in fi_colour_list])
                fi_interpretation_list = \
                                zip(fi_interpretation_list, [self.cell_style(bg=colour) for colour in fi_colour_list])
                row_num_previous = row_num
                row_num = self.write_columns(row_num, 0, [fi_values_list, fi_interpretation_list],
                                                                                                 sheet, columns_length=[1,3])

                di_values_list = ['DI', '1', '>=0.3 but not too close to 1.0', '0.2-0.29', '<=0.19', '0', '-1']
                di_interpretation_list = ['Interpretation',
                                                                  'All HA answered correctly and all LA answered wrongly. Unusual, not desired',
                                                                  'Satisfactory discrimination',
                                                                  'Marginal questions and need some revision',
                                                                  'Poor questions and need major revision or should be eliminated',
                                                                  'Both HA & LA answered correctly. No discrimination at all',
                                                                  'All HA answered wrongly and all LA answered correctly. Problematic item.']
                di_colour_list = ['light_dark_green', 'white', 'light_dark_green', 'light_green', 'light_orange', 'white', 'white']
                di_values_list = zip(di_values_list, [self.cell_style(bg=colour) for colour in di_colour_list])
                di_interpretation_list = \
                                zip(di_interpretation_list, [self.cell_style(bg=colour) for colour in di_colour_list])
                row_num = self.write_columns(row_num_previous, 6, [di_values_list, di_interpretation_list], sheet, columns_length=[3,8])

                # Populate table containing interpretations of both FI and DI at the same time
                fi_values_list = ['FI', '0.3-0.8', '0.20-0.30', '0.00-0.20', '0.80-1.00', 'any value', 'any value']
                di_values_list = ['DI', '>=0.3', '>=0.3', '0.70-0.80', '<0.30', 'negative', 'Zero or close to zero']
                fidi_interpretation_list = ['Interpretation',
                                                                        'Good questions; desired for most of the test questions',
                                                                        'Very difficult questions but can discriminate in a satisfactory manner; use sparingly',
                                                                        'Very difficult questions; testing obscure/specialised knowledge; not recommended',
                                                                        'Easy questions; useful for first few questions to check basic knowledge and skills',
                                                                        'There might be a mistake in the mark scheme; if there is no such error, do not use \
                                                                                        such item',
                                                                        'Do not use such questions as they do not discriminate between low and high achievers']
                fi_values_list = zip(fi_values_list, [self.cell_style(bg='light_green')] * len(fi_values_list))
                di_values_list = zip(di_values_list, [self.cell_style(bg='light_dark_green')] * len(di_values_list))
                fidi_interpretation_list = \
                                zip(fidi_interpretation_list, [self.cell_style(bg='white')] * len(fidi_interpretation_list))
                row_num += 2
                row_num = self.write_columns(row_num, 0, [fi_values_list, di_values_list, fidi_interpretation_list],
                                                                         sheet, columns_length=[1,2,8])

        def write_top_bottom_students(self, top_students, bottom_students, table_header, table_header_style,
                                                                  headers, group_name, sheet):
                '''Write top and bottom students data'''
                column_num = 0
                header_style = self.cell_style(bold=True, black_border=False)
                top_bottom_students_groups = [students[group_name] for students in (top_students, bottom_students)]
                for index, group in enumerate(top_bottom_students_groups):
                        sheet.write(0, column_num, headers[index], header_style)
                        group_styled = []
                        for student in group:
                                group_styled.append(zip(student, table_header_style))
                        self.write_row(1, column_num, zip(table_header, table_header_style), sheet, columns_length=[1,1,3,1])
                        self.write_rows(2, column_num, group_styled, sheet, columns_length=[1,1,3,1])
                        column_num += 7

        def cell_format(self, row_num, column_num):
                # Number of 'A's to insert at the beginning
                column_num_quotient = column_num / 26
                # Can be 'A' to 'Z'
                # Add 1 as row and column numbers are 0-indexed
                column_num_remainder = column_num % 26 + 1
                column_string = 'A'*column_num_quotient+ ('' if column_num_remainder == 0 else chr(column_num_remainder + 64))
                return '{}{}'.format(column_string, row_num+1)

        def write_mcq_students(self, mcq_answers, mcq_students, students_table_header, students_table_header_style, sheet):
                '''Write the students selecting each option for each MCQ question into the sheet'''
                row_num = column_num = 0
                # Used to create the hyperlinks
                cells_all = []
                options = 'ABCD'
                question_style = self.cell_style(bold=True, black_border=False)
                for index, question_students in enumerate(mcq_students):
                        sheet.write_merge(row_num, row_num, column_num, column_num+1,
                                                          'Q{} (Correct: {})'.format(index+1, mcq_answers[index]), question_style)
                        cells_question = []
                        # Since the options are each displayed horizontally, use the option table with
                        # the most rows and base the next question's row number on the last row of that table
                        row_num_question_max = row_num + 2
                        for option_num, option_students in enumerate(question_students):
                                sheet.write(row_num+1, column_num, 'Option {}'.format(options[option_num]))
                                option_students_styled = []
                                for student in option_students:
                                        option_students_styled.append(zip(student, students_table_header_style))
                                if len(option_students_styled) > 0:
                                        self.write_row(row_num+2, column_num, zip(students_table_header, students_table_header_style), sheet,
                                                                   columns_length=[1,1,3])
                                else:
                                        sheet.write(row_num+2, column_num, 'NIL')
                                row_num_option = self.write_rows(row_num+3, column_num, option_students_styled, sheet, columns_length=[1,1,3])
                                row_num_question_max = max(row_num_option, row_num_question_max)
                                cells_question.append(self.cell_format(row_num, column_num))
                                column_num += 6
                        column_num = 0
                        row_num = row_num_question_max + 1
                        cells_all.append(cells_question)
                # The styled MCQ percentages list is a 4x<no. of questions> list while cells_all is a <no. of questions>x4 list
                return zip(*cells_all)

        def write_below_fi_students(self, below_fi_students, students_table_header, students_table_header_style, sheet):
                '''Write the students scoring below FI*weightage for non-MCQ questions'''
                row_num = 0
                cells_all = []
                question_style = self.cell_style(bold=True, black_border=False)
                for index, question_students in enumerate(below_fi_students):
                        sheet.write_merge(row_num, row_num, 0, 10, non_mcq_question_names[index], question_style)
                        cells_all.append(self.cell_format(row_num, 0))
                        row_num += 1
                        question_students_styled = []
                        for student in question_students:
                                question_students_styled.append(zip(student, students_table_header_style))
                        if len(question_students_styled) > 0:
                                row_num = self.write_rows(row_num, 0, question_students_styled, sheet, columns_length=[1,1,3])
                        else:
                                sheet.write(row_num, 0, 'NIL')
                        row_num += 2
                return cells_all

        def write_row(self, row_num, column_num, data, sheet, styled=True, columns_length=None):
                '''Write a 1D array horizontally to an Excel sheet'''
                data = [[item for item in data]]
                return self.write_rows(row_num, column_num, data, sheet, styled, columns_length)

        def write_column(self, row_num, column_num, data, sheet, styled=True, columns_length=None):
                '''Write a 1D array vertically to an Excel sheet'''
                data = [[item for item in data]]
                return self.write_columns(row_num, column_num, data, sheet, styled, columns_length)

        def write_rows(self, row_num, column_num, data, sheet, styled=True, columns_length=None):
                '''Write each subarray in a 2D array as a row to an Excel sheet'''
                column_num_original = column_num
                for row in data:
                        if styled:
                                for index, (item, style) in enumerate(row):                                        
                                        if type(item) == np.int64:
                                                item = int(item)
                                        column_length = (1 if columns_length == None else columns_length[index])
                                        if style == None:
                                                sheet.write_merge(row_num, row_num, column_num, column_num+column_length-1, item)
                                        else:
                                                sheet.write_merge(row_num, row_num, column_num, column_num+column_length-1, item, style)
                                        column_num += column_length
                        else:
                                for index, item in enumerate(row):
                                        column_length = (1 if columns_length == None else columns_length[index])
                                        sheet.write_merge(row_num, row_num, column_num, column_num+column_length-1, item)
                                        column_num += column_length
                        column_num = column_num_original
                        row_num += 1
                return row_num
        
        def write_columns(self, row_num, column_num, data, sheet, styled=True, columns_length=None):
                '''Write each subarray in a 2D array as a column to an Excel sheet'''
                # zip(*data) transposes data
                return self.write_rows(row_num, column_num, zip(*data), sheet, styled, columns_length)

        def round_off(self, val, n):
                '''Round off val to n decimal points'''
                # Use old-style formatting as new-style formatting doesn't seem to be supported
                # in python 2
                if val is not None:
                        return float("%.{}f".format(n) % val)

        def transpose_series(self, vals):
                return vals.to_frame().T

        def percentage(self, a, b):
                if b != 0:
                        return (a+0.0)/b*100

        def mcq_percentage(self, responses, option):
                '''Calculate the percentage of students who chose option'''
                try:
                        mcq_count = self.transpose_series(responses.value_counts())[option]
                except KeyError:
                        # Nobody selected the option
                        return 0
                return self.round_off(self.percentage(mcq_count, len(responses)), 2)

        def process_mcq_responses(self, group_name, responses):
                '''
                Calculate percentage of students selecting each option for each question as well as students who
                selected each option for each question
                '''
                # List of students who chose each option for each question
                mcq_students_all = []
                mcq_percentages_group_index = ['{}_{}'.format(group_name, option) for option in 'ABCD']
                # Stores the percentages of students in the particular selecting each option for all the questions
                # Setting the index and column labels automatically sets the DataFrame's shape, so we can simply
                # set its values by column
                mcq_percentages_group = pd.DataFrame(index=mcq_percentages_group_index,
                                                                                                        columns=mcq_question_numbers)
                options = 'ABCD'
                # [3:] excludes the class, register number and index number
                for column in responses.columns[3:]:
                        # This operation is very time-consuming
                        mcq_students_question = [[], [], [], []]
                        for question_number, option in enumerate(responses[column]):
                                try:
                                        # [0, 1, 2] includes the class, register number and name
                                        option_num = options.find(option)
                                        mcq_students_question[option_num].append(responses.iloc[question_number, range(3)].tolist())
                                except KeyError:
                                        # Empty response
                                        pass
                        mcq_students_all.append(mcq_students_question)
                        
                        # The percentages for the particular question
                        mcq_percentages = []
                        for option in 'ABCD':
                                mcq_percentages.append(self.mcq_percentage(responses[column], option))
                        mcq_percentages_group[column] = mcq_percentages
                return mcq_percentages_group, mcq_students_all

        def fi_students(self, students, weightages):
                '''Return a df containing just the fi values for each question'''
                def fi(vals, weightage):
                        return self.round_off((np.sum(vals)+0.0)/len(vals)/weightage, 2)
                fis = []
                for index, column in enumerate(students.columns):
                        fis.append(fi(students[column], weightages[index]))
                return pd.Series(fis)

        def process_fidi(self, group_name, results, weightages):
                '''Return a Series of FI and DI values for each question'''
                top_bottom_cutoff = int(0.27 * len(results.index))
                top_students = results.sort_values('total_score', axis=0, ascending=False)[:top_bottom_cutoff]
                bottom_students = \
                        results.sort_values('total_score', axis=0, ascending=True)[:top_bottom_cutoff]
                # Ignore the class, index number, name and total score when calculating the df
                columns_questions_results = range(3, len(results.columns)-1)
                fi_df = self.fi_students(results.iloc[:, columns_questions_results], weightages)
                di_df = self.fi_students(top_students.iloc[:, columns_questions_results], weightages) - \
                                self.fi_students(bottom_students.iloc[:, columns_questions_results], weightages)
                # "Transpose" the series
                fi_df = self.transpose_series(fi_df); di_df = self.transpose_series(di_df)
                # Change the name so that it will be reflected in fidi_df
                fi_df = fi_df.rename(lambda x: group_name+'_fi')
                di_df = di_df.rename(lambda x: group_name+'_di')
                # Only retrieve the class, index number, name and total score of the student
                # Convert to 2D array for easier iteration later on
                columns_no_questions_results = [0, 1, 2, len(top_students.columns)-1]
                top_students_no_questions_results = top_students.iloc[:, columns_no_questions_results].values.tolist()
                bottom_students_no_questions_results= bottom_students.iloc[:, columns_no_questions_results].values.tolist()
                # Return top_students and bottom_students for data drill down
                return fi_df, di_df, top_students_no_questions_results, bottom_students_no_questions_results

        def process_below_fi_students(self, results_df, non_mcq_fidi_df, group_names):
                below_fi_students_all = {}
                # Keep only the class, register number, name and non-mcq results
                results_df_non_mcq_list = results_df.drop(results_df.columns[3:3+mcq_question_length], axis=1).values.tolist()
                if include_non_mcq:
                        for group_name in group_names:
                                below_fi_students_group = []
                                for question_num in range(non_mcq_question_length):
                                        below_fi_students_question = []
                                        fi_threshold = \
                                non_mcq_question_weightages[question_num] * non_mcq_fidi_df.loc['{}_fi'.format(group_name)] \
                                                                                                                        [(25+question_num if include_mcq else question_num)]
                                        students_group = []
                                        if group_name == 'batch':
                                                # Just include the whole list
                                                students_group = results_df_non_mcq_list
                                        else:
                                                for student in results_df_non_mcq_list:
                                                        if student[0] == group_name:
                                                                students_group.append(student)
                                        for student in students_group:
                                                if student[3+question_num] < fi_threshold:
                                                        below_fi_students_question.append(student)
                                        below_fi_students_group.append(below_fi_students_question)
                                below_fi_students_all.update({group_name: below_fi_students_group})
                return below_fi_students_all

        def cell_style(self, fg='black', bg='white', black_border=True, bold=False, wrap=False, align_horizontal='center', align_vertical='center'):
                '''Return the style which sets the foreground and background colour with optional black border'''
                bold_style_string = 'bold {};'.format(bold)
                border_style_string = 'borders: top_color black, bottom_color black, right_color black, left_color black'
                if black_border:
                        border_style_string += ', left thin, right thin, top thin, bottom thin;'
                else:
                        border_style_string += ';'
                align_style_string = 'align: wrap {}, horizontal {}, vertical {}'.format(wrap, align_horizontal, align_vertical)
                background_style_string = 'pattern: pattern solid, fore_colour {};'.format(bg)
                foreground_style_string = 'font: colour {}'.format(fg)
                return xlwt.Style.easyxf('{} {}, {} {} {}'.format(background_style_string, foreground_style_string, bold_style_string,
                                                                                                                  border_style_string, align_style_string))

        def corresponding_style(self, map, value):
                '''Return the style based on a map acting on a value'''
                for condition, style in map.iteritems():
                        if condition(value):
                                return style
                return self.cell_style()

        def create_section_template(self, section_name, question_names, fis, dis, document):
                '''Create the table template for the teachers to fill in the report'''
                table = document.add_table(rows=2, cols=0)
                table.style = 'Table Grid'
                # The only way to set the column widths that works
                table.add_column(Cm(COLUMN_QUESTION_WIDTH))
                table.add_column(Cm(COLUMN_TOPIC_WIDTH))
                table.add_column(Cm(COLUMN_MARKER_WIDTH))
                table.add_column(Cm(COLUMN_COMMENTS_WIDTH))
                section_heading = table.cell(0, 0)
                section_heading.text = section_name
                section_heading.merge(table.cell(0, 1))
                section_heading.merge(table.cell(0, 2))
                section_heading.merge(table.cell(0, 3))
                column_question = table.cell(1, 0)
                column_question.width = Cm(0.5)
                column_topic = table.cell(1, 1)
                column_marker = table.cell(1, 2)
                column_comments = table.cell(1, 3)
                column_question.text = 'Qn.'
                column_topic.text = 'Topic'
                column_marker.text = 'Marker'
                column_comments.text = 'Comments'
                # Since there is no official way to set a background colour of the cell, inject the XML directly
                # Create new elements for each cell or else only the last cell will have the style
                gray_background_xml_1 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), GRAY))
                gray_background_xml_2 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), GRAY))
                gray_background_xml_3 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), GRAY))
                gray_background_xml_4 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), GRAY))
                column_question._tc.get_or_add_tcPr().append(gray_background_xml_1)
                column_topic._tc.get_or_add_tcPr().append(gray_background_xml_2)
                column_marker._tc.get_or_add_tcPr().append(gray_background_xml_3)
                column_comments._tc.get_or_add_tcPr().append(gray_background_xml_4)
                for index in range(len(fis)):
                        row = table.add_row()
                        row.cells[0].width = Cm(0.5)
                        row.cells[0].text = question_names[index]
                        paragraph_fidi = row.cells[0].add_paragraph()
                        paragraph_fidi.text = 'FI: {}'.format(fis[index])
                        paragraph_fidi_spacer = paragraph_fidi.add_run()
                        paragraph_fidi_spacer.add_break()
                        paragraph_fidi.text += 'DI: {}'.format(dis[index])

        def calculate_questions_difficulty(self, fis):
                '''Calculate the percentage of questions for each difficulty from the FI value'''
                easy_count, moderate_count, difficult_count = 0, 0, 0
                for fi in fis:
                        if fi < 0.4:
                                difficult_count += 1
                        elif 0.4 <= fi <= 0.69:
                                moderate_count += 1
                        else:
                                easy_count += 1
                count_percentages = []
                for count in easy_count, moderate_count, difficult_count:
                        count_percentages.append(self.round_off(self.percentage(count, len(fis)), 2))
                return count_percentages

if __name__ == '__main__':
        app = wx.App(False)
        frame = MainFrame(parent=None)
        frame.Show()
        app.MainLoop()
